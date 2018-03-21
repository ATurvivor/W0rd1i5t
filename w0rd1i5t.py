# -*- coding: utf-8 -*-
"""
Created on Wed Mar 21 01:07:42 2018

@author: fqy

w0rd1ist.py: g3n3rat3 w0rd1ist fr0m d0cx fi135 :>
"""

import docx
import regex as re
import string
from os import listdir
from os.path import isfile, join
from reportlab.pdfgen import canvas

class verbum:
    
    def __init__(self):
    #initialize class verbum 
        self.word = ""
        self.source = ""
        
    def __lt__(self, other):
    #sorting priority: less than
        return self.word < other.word
        
#regular expression pattern
pattern_zh = r'([\p{IsHan}\p{IsBopo}\p{IsHira}\p{IsKatakana}]+)'
pattern_phon = r'[\/|\[].*[\/\]]'
pattern_word = r'\w+\s+'

def s3p_zh(input):
    # s3parat3 chin353 charact3r5 with @...@
    pattern = re.compile(pattern_zh, re.UNICODE)
    output = pattern.sub(r'@\1@', input)
    return output

def i5_h3ad(input):
    # h3ad 3ntry 0r n0t->
        # -zh_m3aning: x
        # phon3tiqu3: x
    # return entry word or null
    searchObj = re.search(pattern_phon, input)
    if searchObj:
            return re.search(pattern_word, input)
    else:
        return searchObj
            
"""
sample_filename = 'C:\\Users\\admin\\Desktop\\dic\\3.9.docx'
sample_document = docx.Document(sample_filename);
sample_wordlist = [];

for p in sample_document.paragraphs:
    h3ad = i5_h3ad(p.text)
    if h3ad:
        word = h3ad.group()
        sample_wordlist.append(word.capitalize())
    
print(sample_wordlist)
"""
w0rd1i5t = [];
filepath_p1 = 'C:\\Users\\admin\\Desktop\\dic\\2017 Vocab Database\\'
_p2 = ['ARQ---Chen Xin', 'BFGPW---Sun Chuyue', 'COHMUN---Chen Chen', 'DEV---Du Weiyu', 'ILT-JK---Zhao Zixuan', 'SP---Chen Xiao']

for i in range(6):
    filepath = filepath_p1+_p2[i];
    fp =[f for f in listdir(filepath) if isfile(join(filepath,f))]
    
    for df in range(len(fp)):
        document = docx.Document(filepath_p1+_p2[i]+'\\'+fp[df]);
    
        for p in document.paragraphs:
            h3ad = i5_h3ad(s3p_zh(p.text))
            if h3ad:
                ws = h3ad.group()
                temp = verbum()
                temp.word = ws.capitalize()
                temp.word = re.sub(r'\s+$', "", temp.word)
                temp.source = _p2[i]+'/'+fp[df]
                #temp.word.replace(u'\xa0', u' ')
                #temp.source.replace(u'\xa0', u' ')
                if temp.word != 'Verb' and temp.word != 'Adjective' and temp.word != 'Of' and temp.word!='Used' and temp.word!='Late' and temp.word!='If' and temp.word!='I' and temp.word!='For' and temp.word!='From' and temp.word!='Early' and temp.word != 'To' and temp.word !='Or' and temp.word != 'One' and temp.word != 'En'  and temp.word != 'British' and temp.word !='American' and temp.word!= 'Any' and temp.word !='1786' and temp.word!='As' and temp.word != 'Noun' and temp.word!='An'and temp.word!='A' and temp.word!= 'Also':
                                    w0rd1i5t.append(temp)

w0rd1i5t.sort()

#print
"""
canvas = canvas.Canvas("C:\\Users\\admin\\Desktop\\dic\\w0rd1i5t.pdf")
canvas.setFont('Helvetica', 20)
canvas.drawString(250,750, "WORDLIST")
canvas.setFont('Helvetica', 15)
for i in range(len(w0rd1i5t)):
    canvas.drawString(30,720-i*20, w0rd1i5t[i].word)
    canvas.drawString(100,720-i*20, w0rd1i5t[i].source)
    
canvas.save();
"""

file = open("gen_latex.txt","w",encoding='utf-8') 
 
for i in range(len(w0rd1i5t)):
    file.write("\hline\n")
    temp_ = w0rd1i5t[i].word+"  &  "+w0rd1i5t[i].source+"\\\  \n"
    file.write(temp_)
 
file.close() 

"""
for i in range(50,100):
    print(w0rd1i5t[i].word)
    print (w0rd1i5t[i].source)
    print
    """